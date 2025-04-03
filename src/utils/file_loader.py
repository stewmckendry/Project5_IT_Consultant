# file_loader.py â€“ File input handling

from pathlib import Path
import docx # for Word documents
import fitz # for PDFs
import os
from typing import Dict
import PyPDF2
from src.utils.rfp_extractors import extract_evaluation_criteria
import re
from src.utils.logging_utils import log_phase

def load_report_text_from_file(filepath=None, file=None) -> str:
    """
    Loads text from a supported file format (txt, md, docx, pdf).
    """
    if filepath:
        ext = Path(filepath).suffix.lower()
    elif file:
        ext = Path(file.filename).suffix.lower()
    else:
        raise ValueError("Must provide either 'filepath' or 'file'")

    if ext in [".txt", ".md"]:
        content = file.file.read().decode("utf-8") if file else Path(filepath).read_text(encoding="utf-8")
    elif ext == ".docx":
        doc = docx.Document(file.file if file else filepath)
        content = "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        doc = fitz.open(stream=file.file.read(), filetype="pdf") if file else fitz.open(filepath)
        content = "\n".join([page.get_text() for page in doc])
    else:
        raise ValueError("Unsupported file format. Use .txt, .md, .docx, or .pdf")
    
    return content


def load_proposals_from_folder(folder_path: str) -> Dict[str, str]:
    """
    Loads all proposal files in a folder (supports .txt, .pdf, .docx) and returns a dict.
    Keys are inferred from filenames (e.g., 'Vendor A.txt' â†’ 'Vendor A').
    """
    proposals = {}
    folder = Path(folder_path)

    for file in folder.glob("*"):
        vendor_name = file.stem  # e.g., "Vendor A"

        if file.suffix == ".txt":
            proposals[vendor_name] = file.read_text(encoding="utf-8")

        elif file.suffix == ".docx":
            doc = docx.Document(file)
            proposals[vendor_name] = "\n".join([para.text for para in doc.paragraphs])

        elif file.suffix == ".pdf":
            with open(file, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
                proposals[vendor_name] = text

    return proposals


def load_rfp_criteria(filepath: str) -> list:
    """
    Loads RFP evaluation criteria from a file.
    Each line should represent one criterion.
    """
    text = load_report_text_from_file(filepath)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    
    # Remove numbering if present
    criteria = [line.split('.', 1)[-1].strip() if '.' in line else line for line in lines]
    return criteria


def parse_rfp_from_file(filepath: str) -> dict:
    """
    Parses RFP file and extracts:
    - Full text
    - Section map (e.g., Background, Requirements, Criteria)
    - Evaluation criteria with weights and descriptions (if found)
    """
    full_text = load_report_text_from_file(filepath)
    
    # Step 1: Split sections using common RFP headers
    section_patterns = {
        "Background": r"(?:^|\n)(Background|Introduction)\s*[\n:-]",
        "Requirements": r"(?:^|\n)(Requirements|Business Needs|Scope of Work)\s*[\n:-]",
        "Evaluation Criteria": r"(?:^|\n)(Evaluation Criteria|Scoring Methodology|Proposal Evaluation)\s*[\n:-]",
        "Terms and Conditions": r"(?:^|\n)(Terms and Conditions|Contract Requirements)\s*[\n:-]",
    }

    sections = {}
    for name, pattern in section_patterns.items():
        match = re.search(pattern, full_text, flags=re.IGNORECASE)
        if match:
            start_idx = match.start()
            sections[name] = full_text[start_idx:]

    # Step 2: Try to extract evaluation criteria block
    criteria_block = sections.get("Evaluation Criteria", "")
    criteria = extract_evaluation_criteria(criteria_block)

    return {
        "full_text": full_text,
        "sections": sections,
        "criteria": criteria
    }


from typing import List, Dict, Optional
from sentence_transformers import SentenceTransformer, util

model = SentenceTransformer("all-MiniLM-L6-v2")

def preprocess_proposal_for_criteria_with_threshold(
    proposal_text: str,
    rfp_criteria: List[str],
    rfp_criterion_descriptions: Optional[Dict[str, str]] = None,
    score_threshold: float = 0.4
) -> Dict[str, str]:
    """
    Matches segments of proposal text to each RFP criterion using embedding similarity,
    returning only segments above a relevance threshold.
    """
    # Split proposal into paragraphs
    paragraphs = [p.strip() for p in proposal_text.split("\n") if p.strip()]
    para_embeddings = model.encode(paragraphs, convert_to_tensor=True)

    matched_sections = {}

    log_phase("ðŸ” Matching proposal sections to RFP criteria...")
    for criterion in rfp_criteria:
        criterion = criterion["name"]
        query = rfp_criterion_descriptions.get(criterion, criterion) if rfp_criterion_descriptions else criterion
        query_embedding = model.encode(query, convert_to_tensor=True)

        # Compute similarity scores
        cosine_scores = util.pytorch_cos_sim(query_embedding, para_embeddings)[0]

        # Select all paragraphs above the threshold
        relevant_paragraphs = [
            para for para, score in zip(paragraphs, cosine_scores)
            if score.item() >= score_threshold
        ]
        log_phase(f"ðŸ” Found {len(relevant_paragraphs)} relevant paragraphs for criterion '{criterion}'")

        # If nothing matched, fall back to top 1 best match
        if not relevant_paragraphs:
            log_phase(f"ðŸ” No paragraphs above threshold for '{criterion}'. Selecting top match.")
            top_idx = cosine_scores.argmax().item()
            relevant_paragraphs = [paragraphs[top_idx]]

        matched_sections[criterion] = "\n\n".join(relevant_paragraphs)

    return matched_sections


def load_scenario_data(scenario_name, base_path: Path = Path("../data/rfp_scenarios")):
    scenario_path = (base_path / scenario_name).resolve()

    if not scenario_path.is_dir():
        raise FileNotFoundError(f"Scenario directory not found: {scenario_path}")

    rfp_file = scenario_path / "rfp.txt"
    log_phase(f"ðŸ“ Loading scenario data from {scenario_path} (RFP: {rfp_file})")

    proposals = {}
    txt_files = list(scenario_path.glob("*.txt"))
    log_phase(f"ðŸ“„ Found {len(txt_files)} files")  # Count of files

    for file in txt_files:
        log_phase(f"ðŸ“„ Loading proposal from {file.name}")
        if file.name.lower() != "rfp.txt":
            vendor = file.stem.replace("_", " ").title()
            proposals[vendor] = file.read_text()
            log_phase(f"ðŸ“„ Loaded proposal for {vendor}")

    return proposals, str(rfp_file)


DEFAULT_SCENARIO = "scenario1_basic"
DEFAULT_SCENARIO_DIR = Path("../data/rfp_scenarios")


def load_default_scenario(scenario_name: str = DEFAULT_SCENARIO) -> tuple[dict, str]:
    """
    Loads RFP and vendor proposals from a named scenario folder.

    Parameters:
        scenario_name (str): Folder name under data/rfp_scenarios (e.g., "basic_test").

    Returns:
        proposals (dict): {vendor_name: proposal_text}
        rfp_path (str): Path to the scenario's RFP file
    """
    log_phase(f"ðŸ“ Loading scenario: {scenario_name}")
    proposals, rfp_file = load_scenario_data(scenario_name, base_path=DEFAULT_SCENARIO_DIR)

    log_phase(f"âœ… Loaded {len(proposals)} proposals and RFP from {rfp_file}")
    return proposals, rfp_file

def list_available_scenarios(base_path: Path = DEFAULT_SCENARIO_DIR):
    return [f.name for f in base_path.iterdir() if f.is_dir()]


from typing import List, Tuple, Dict
from fastapi import UploadFile
import tempfile
from pathlib import Path

def process_uploaded_files(files: List[UploadFile]) -> Tuple[Dict[str, str], str]:
    """
    Processes uploaded files and separates RFP from proposals.

    Returns:
        proposals (dict): {vendor_name: proposal_text}
        rfp_path (str): Path to saved RFP file
    """
    temp_dir = Path(tempfile.mkdtemp(prefix="rfp_eval_"))
    proposals = {}
    rfp_path = None

    for file in files:
        contents = load_report_text_from_file(file=file)
        filename = Path(file.filename).name.lower()

        if "rfp" in filename:
            rfp_path = temp_dir / "rfp.txt"
            rfp_path.write_text(contents)
        else:
            vendor_name = Path(file.filename).stem.replace("_", " ").title()
            proposal_path = temp_dir / file.filename
            proposal_path.write_text(contents)
            proposals[vendor_name] = contents
        
        log_phase(f"ðŸ“Ž Loaded {file.filename} â†’ {'RFP' if 'rfp' in filename else 'Proposal'}")

    if not rfp_path or not proposals:
        raise ValueError("Missing RFP file or no proposals uploaded.")

    return proposals, str(rfp_path)
